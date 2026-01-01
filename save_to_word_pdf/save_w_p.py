from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import arabic_reshaper
from bidi.algorithm import get_display

from django.contrib import messages
from django.http import FileResponse
from io import BytesIO

out_dir = "media"
filename = "my_report"

# test_text = "تست اول test 2 تست سوم"
test_text = """
    دیروز به پارک hhh رفتم. به جان دوتایی‌ام قسم ،اگر دروغ بگویم، مردی داشت از درخت بالا می‌رفت. اصلاً باورم نمی‌شد؛ من کاملاً هاج و واج مانده بودم. مگر می‌شود؟ مگر داریم؟ آن مرد رفته بود دنبال گربه. گربه هم مدام از آن بالا میو میو می‌کرد. من نمی‌دانم، آخه برای چه این کار را می‌کرد؟ گربه به تو چه ربطی دارد؟ حالا ما مدام می‌گفتیم: «ای احمق، به خدا بیا پایین؛ گربه بی‌گناه است، زبون‌بسته خداست، خدا لعنتت کند.» بعدش رفتیم نزد نگهبان پارک و گفتیم، اما صدای گربه دیگر در نمی‌آمد. ما هیچ به هیچ؛ هرچه ما می‌گفتیم، هیچ‌کس گوش نمی‌داد؛ نه نگهبان توجه می‌کرد، نه مردم کاری داشتند. اصلاً گربه، بنده خدا، دیگر اصلاً صدایی از او در نمی‌آمد. دیگر رفتیم چوبی برداشتیم و افتادیم به جان آن مرد بالای درخت؛ می‌گفتیم: «بیا پایین، بیا پایین؛ تا نکشمت هیچی.» هیچی، باورم نمی‌شد که اوضاع این‌طوری باشد. چوب را پرت کردیم، زدیم به سرش، زدیم به دست و پایش؛ هیچی. می‌گفتیم: «بابا بیا بیرون، بیا پایین، بیا پایین.» که گوش می‌داد، آره. خلاصه.
    """

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

FONT_PATH = os.path.join(BASE_DIR, "Vazir-Regular.ttf")
FONT_NAME = "Vazir"

pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))


def save_as_word(text, title="test"):
    try:
        document = Document()
        style = document.styles['Normal']

        font = style.font
        font.size = Pt(12)

        style.paragraph_format.bidi = True

        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_heading(title)

        # مثال: اضافه کردن پاراگراف‌ها
        for paragraph_text in text.split('\n'):
            if paragraph_text.strip():
                paragraph = document.add_paragraph(paragraph_text, style=style)
                # تنظیم جهت متن (برای فارسی: راست به چپ)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                paragraph.paragraph_format.bidi = True
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # ----------------------------

        # ذخیره فایل در حافظه موقت (BytesIO)
        word_buffer = BytesIO()
        document.save(word_buffer)
        word_buffer.seek(0)  # بازگرداندن نشانگر به ابتدا

        # بازگرداندن آبجکت BytesIO که حاوی داده‌های فایل Word است
        return word_buffer




    except Exception as e:
        print("ssssss")
        return f"خطا در ایجاد فایل Word: {e}"


try:
    pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))
    # print(f"✅ فونت {FONT_NAME} با موفقیت ثبت شد.")
except Exception as e:
    # اگر فایل فونت Vazir پیدا نشد، خطا می‌دهد و از Helvetica استفاده می‌کند
    # print(
    #     f"❌ خطای حیاتی فونت: فایل {FONT_PATH} یافت نشد. لطفا فایل فونت را در مسیر اجرا قرار دهید تا مشکل مربع‌ها حل شود. جزئیات خطا: {e}")
    FONT_NAME = "Helvetica"  # بازگشت به فونت پیش‌فرض (اما با مربع)

def wrap_line(text, font_name, font_size, max_width):
    words = text.split(' ')
    lines = []
    current_line = ""
    for word in words:
        test_line = current_line + (" " if current_line else "") + word
        reshaped = arabic_reshaper.reshape(test_line)
        bidi_line = get_display(reshaped)
        width = pdfmetrics.stringWidth(bidi_line, font_name, font_size)
        if width > max_width and current_line:
            lines.append(current_line)
            current_line = word
        else:
            current_line = test_line
    if current_line:
        lines.append(current_line)
    return lines

def save_as_pdf(text):
    pdf_buffer = BytesIO()

    c = canvas.Canvas(pdf_buffer, pagesize=A4)
    font_name = "Vazir"
    font_size = 14
    c.setFont(font_name, font_size)
    width, height = A4
    margin = 50
    y = height - margin

    # برای هر خط از متن اصلی
    for paragraph in text.split('\n'):
        # پاراگراف را به خطوط کوتاه‌تر تقسیم کن
        wrapped_lines = wrap_line(paragraph, font_name, font_size, width - 2*margin)
        for line in wrapped_lines:
            reshaped_text = arabic_reshaper.reshape(line)
            bidi_text = get_display(reshaped_text)
            # print(bidi_text)
            c.drawRightString(width - margin, y, bidi_text)
            y -= font_size + 6  # فاصله بین خطوط
            if y < margin:  # اگر به پایین صفحه رسیدیم صفحه جدید بساز
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - margin
    c.showPage()
    c.save()

    # بازگرداندن نشانگر به ابتدای بافر برای خواندن توسط جنگو
    pdf_buffer.seek(0)

    # بازگرداندن آبجکت BytesIO
    return pdf_buffer
